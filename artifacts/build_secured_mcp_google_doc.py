from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "secured-mcp-registry-user-journeys.docx"
DIAGRAM = ROOT / "secured-mcp-business-journey.png"
PUBLISHER_SEQUENCE = ROOT / "publisher-journey-sequence.png"
AGENT_SEQUENCE = ROOT / "ai-agent-journey-sequence.png"


def set_font(run, size=11, bold=False, color="000000"):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=8, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_body(doc, text, *, bold_lead=None):
    p = doc.add_paragraph()
    set_spacing(p)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead) :])
        set_font(rest)
    else:
        run = p.add_run(text)
        set_font(run)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    sizes = {1: 20, 2: 16, 3: 14}
    before = {1: 20, 2: 18, 3: 16}
    after = {1: 6, 2: 6, 3: 4}
    set_spacing(p, before=before[level], after=after[level])
    run = p.add_run(text)
    set_font(run, size=sizes[level], color="000000")
    return p


def add_number(doc, title, detail):
    p = doc.add_paragraph(style="List Number")
    set_spacing(p, after=6)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_font(p.add_run(title + " — "), bold=True)
    set_font(p.add_run(detail))
    return p


def add_numbered_statement(doc, title, detail):
    return add_number(doc, title, detail)


def add_figure(doc, path, caption_text, width=6.45):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=8, after=4)
    p.add_run().add_picture(str(path), width=Inches(width))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(caption, after=12)
    set_font(caption.add_run(caption_text), size=10, color="555555")


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

for level, size, before, after, color in [
    (1, 20, 20, 6, "000000"),
    (2, 16, 18, 6, "000000"),
    (3, 14, 16, 4, "434343"),
]:
    style = styles[f"Heading {level}"]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(size)
    style.font.bold = False
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

title = doc.add_paragraph()
set_spacing(title, after=3)
set_font(title.add_run("Secured MCP Registry"), size=26)

subtitle = doc.add_paragraph()
set_spacing(subtitle, after=18)
set_font(
    subtitle.add_run("Publisher and AI Agent User Journeys"),
    size=14,
    color="555555",
)

add_body(
    doc,
    "A simple guide to how MCP services are published, discovered and used safely by AI agents and LLM-powered applications.",
)

add_heading(doc, "At a glance", 1)
add_body(
    doc,
    "The Secured MCP Registry is a trusted directory for MCP services. Publishers use it to make services available, while AI agents use it to find services that have passed identity, safety and permission checks.",
)
add_body(
    doc,
    "The registry is more than a catalogue. It continues to watch for risk and can restrict or stop access when behaviour changes.",
)

add_figure(
    doc,
    DIAGRAM,
    "How publishers, the registry and AI agents work together",
)

add_heading(doc, "Who takes part?", 1)
add_numbered_statement(
    doc,
    "Publisher",
    "The person or organization that creates and maintains an MCP service.",
)
add_numbered_statement(
    doc,
    "Secured MCP Registry",
    "The trusted platform that verifies, approves, publishes and monitors MCP services.",
)
add_numbered_statement(
    doc,
    "AI Agent / LLM Application",
    "The user-facing application that decides when an MCP capability is useful.",
)
add_numbered_statement(
    doc,
    "Agent Host / MCP Client",
    "The software that connects the AI agent to an approved MCP service.",
)

add_heading(doc, "Publisher journey", 1)
add_body(
    doc,
    "The publisher journey is designed to make useful MCP services available without asking users to accept unknown or unnecessary risk.",
)
add_number(
    doc,
    "Create",
    "The publisher builds an MCP service and explains what it does.",
)
add_number(
    doc,
    "Describe",
    "The publisher provides ownership, capability, permission and version information.",
)
add_number(
    doc,
    "Submit",
    "The service is sent to the Secured MCP Registry for review.",
)
add_number(
    doc,
    "Verify",
    "The registry checks identity, safety information, permissions and the service’s signed proof.",
)
add_number(
    doc,
    "Approve and publish",
    "An approved service becomes discoverable to eligible AI agents and applications.",
)
add_number(
    doc,
    "Maintain",
    "The publisher releases updates and responds to any new safety findings.",
)
add_figure(
    doc,
    PUBLISHER_SEQUENCE,
    "Numbered publisher journey through the Secured MCP Registry",
)

add_heading(doc, "AI Agent / LLM Application journey", 1)
add_body(
    doc,
    "“Consumer” is the business role. In practice, the consumer is usually an AI agent or LLM-powered application operating through an MCP client.",
)
add_number(
    doc,
    "Understand the need",
    "The AI agent identifies a capability required to complete the user’s task.",
)
add_number(
    doc,
    "Discover",
    "The MCP client searches the Secured MCP Registry for suitable approved services.",
)
add_number(
    doc,
    "Check",
    "The registry confirms the publisher, version, safety status and requested permissions.",
)
add_number(
    doc,
    "Authorize",
    "Access is allowed only when organizational policy and user permission requirements are met.",
)
add_number(
    doc,
    "Use",
    "The AI agent invokes the approved MCP service and receives the result.",
)
add_number(
    doc,
    "Record",
    "The important decisions and actions are recorded for accountability.",
)
add_number(
    doc,
    "Respond to change",
    "Access can continue, require approval, be restricted or be stopped when risk changes.",
)
add_figure(
    doc,
    AGENT_SEQUENCE,
    "Numbered AI agent and LLM application journey",
)

add_heading(doc, "The pillars of SecureMCP", 1)
add_body(
    doc,
    "SecureMCP uses several connected safeguards rather than relying on a single security check. Each pillar answers a different business question, from “Can we trust this service?” to “Should this action continue right now?”",
)

pillars = [
    (
        "1. Secured MCP Registry",
        "The trusted system of record for MCP publishers, services, versions and security status.",
        "It gives organizations one governed place to publish, discover, approve, suspend and retire MCP services.",
    ),
    (
        "2. Tool Marketplace and Discovery",
        "The approved catalogue through which AI agents and applications find suitable MCP capabilities.",
        "It makes discovery useful without treating every available service as automatically trusted.",
    ),
    (
        "3. Certification and Attestation",
        "The evidence that a publisher and service completed the required checks for a particular version.",
        "Signed proof allows an organization to verify what was assessed and whether that proof is still valid.",
    ),
    (
        "4. Policy",
        "The organization’s rules for who may do what, under which conditions and within which limits.",
        "Policy decisions can allow an action, deny it or require additional approval.",
    ),
    (
        "5. Contracts",
        "A clear usage agreement between the AI agent and the MCP service.",
        "It defines purpose, permitted scope, obligations and limits before a capability is used.",
    ),
    (
        "6. Consent",
        "The record of permission from the appropriate user, owner or organization.",
        "Consent ensures that technically possible actions are not performed without appropriate authority.",
    ),
    (
        "7. Provenance",
        "A traceable history of where a request came from, which decisions were made and what happened.",
        "Provenance supports accountability, investigation and confidence in the final outcome.",
    ),
    (
        "8. Reflexive Core",
        "The adaptive safety layer that observes behaviour, detects drift and responds to changing risk.",
        "It can allow normal activity, request approval, restrict access or stop a critical action.",
    ),
    (
        "9. Introspection",
        "The ability to examine the current request, actor, context and risk before or during execution.",
        "Introspection gives the Reflexive Core the context required to make a meaningful assessment.",
    ),
    (
        "10. Alerts",
        "The shared security event and notification capability.",
        "Alerts bring unusual activity, policy failures and urgent risks to the attention of responsible people and systems.",
    ),
    (
        "11. Gateway and Audit",
        "The governed entry point for security-aware MCP access and a unified view of security activity.",
        "It simplifies oversight by bringing access and audit information into one controlled layer.",
    ),
    (
        "12. Compliance",
        "The reporting layer that turns security evidence into information suitable for governance and assurance.",
        "It helps demonstrate that required controls are operating and provides material for reviews or audits.",
    ),
    (
        "13. Federation",
        "A way for trusted SecureMCP environments to share approved trust information.",
        "Federation supports collaboration across teams or organizations without removing local control.",
    ),
    (
        "14. Certificate Revocation",
        "The mechanism for declaring that previously trusted proof must no longer be accepted.",
        "Revocation protects users when a credential, publisher or service version becomes unsafe or compromised.",
    ),
    (
        "15. Sandboxed Execution",
        "An isolated environment for running higher-risk capabilities with controlled access to systems and data.",
        "Sandboxing limits the impact of unexpected or unsafe behaviour.",
    ),
]

for title_text, meaning, value in pillars:
    heading = add_heading(doc, title_text, 2)
    keep_with_next(heading)
    add_body(doc, f"What it is: {meaning}", bold_lead="What it is:")
    add_body(doc, f"Why it matters: {value}", bold_lead="Why it matters:")

add_heading(doc, "Platform foundations around the pillars", 1)
add_body(
    doc,
    "The PureCipher Secured MCP Registry adds the practical product services needed to operate these pillars:",
)
foundations = [
    (
        "Identity and role-based access",
        "confirms publishers, administrators, reviewers and registered MCP clients.",
    ),
    (
        "Publisher and submission services",
        "support packaging, preflight checks, submission and version updates.",
    ),
    (
        "Moderation and lifecycle management",
        "support approval, rejection, suspension, deprecation and deregistration.",
    ),
    (
        "Client identity and tokens",
        "give each AI agent, service or framework a controlled identity.",
    ),
    (
        "OpenAPI ingestion and MCP gateway",
        "allow eligible API operations to be presented through governed MCP toolsets.",
    ),
    (
        "Notifications",
        "inform publishers, reviewers, administrators and client owners when action is required.",
    ),
    (
        "Durable records",
        "retain registry, account, security, client and audit information.",
    ),
]
for name, detail in foundations:
    add_body(doc, f"{name}: {detail}", bold_lead=f"{name}:")

add_heading(doc, "How the pillars work together", 1)
for name, detail in [
    (
        "Establish trust",
        "Identity, the Secured MCP Registry, certification and revocation establish whether a service should be considered trustworthy.",
    ),
    (
        "Set boundaries",
        "Policy, contracts and consent establish what the AI agent is allowed to do.",
    ),
    (
        "Execute safely",
        "The gateway, sandbox and MCP server apply those boundaries during use.",
    ),
    (
        "Observe and adapt",
        "Introspection and the Reflexive Core detect changing behaviour and adjust the response.",
    ),
    (
        "Maintain accountability",
        "Provenance, audit, alerts and compliance preserve evidence and bring important events to attention.",
    ),
    (
        "Share trust carefully",
        "Federation allows trusted environments to collaborate while retaining local governance.",
    ),
]:
    add_body(doc, f"{name}: {detail}", bold_lead=f"{name}:")

add_heading(doc, "What the audience should remember", 1)
add_numbered_statement(
    doc,
    "Publisher assurance",
    "Publishers have a clear path to prove that their MCP services are trustworthy.",
)
add_numbered_statement(
    doc,
    "Controlled discovery",
    "AI agents discover and use approved services through controlled access.",
)
add_numbered_statement(
    doc,
    "Continuous safety",
    "Safety is checked continuously, not only at publication time.",
)
add_numbered_statement(
    doc,
    "Adaptive response",
    "The Reflexive Core can adapt the response when behaviour or risk changes.",
)
add_numbered_statement(
    doc,
    "Accountability",
    "Every important action can be traced for review and accountability.",
)

props = doc.core_properties
props.title = "Secured MCP Registry — Publisher and AI Agent User Journeys"
props.subject = "Non-technical overview of Secured MCP Registry user journeys"
props.author = "PureCipher"

doc.save(OUTPUT)
print(OUTPUT)
