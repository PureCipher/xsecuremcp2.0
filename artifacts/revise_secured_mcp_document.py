from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SOURCE = Path("/Users/purecipher/Downloads/Secured MCP Registry.docx")
OUTPUT = Path(
    "/Users/purecipher/code/xsecuremcp2.0/artifacts/Secured MCP Registry - Executive Overview.docx"
)
FLOW = Path(
    "/Users/purecipher/code/xsecuremcp2.0/artifacts/secured-mcp-high-level-flow.png"
)


def set_arial(run, size=None, bold=None, color=None):
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Arial")
    rpr.rFonts.set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def move_after(paragraph, anchor_element):
    element = paragraph._element
    anchor_element.addnext(element)
    return element


def add_body(doc, text, *, lead=None):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    if lead and text.startswith(lead):
        r1 = p.add_run(lead)
        set_arial(r1, size=11, bold=True)
        r2 = p.add_run(text[len(lead) :])
        set_arial(r2, size=11)
    else:
        set_arial(p.add_run(text), size=11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_arial(p.add_run(text))
    return p


def add_numbered_item(doc, title, detail):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    set_arial(p.add_run(title + " — "), size=11, bold=True)
    set_arial(p.add_run(detail), size=11)
    return p


doc = Document(SOURCE)
original = list(doc.paragraphs)

# Update the subtitle while retaining its original paragraph formatting.
subtitle = original[1]
subtitle.clear()
set_arial(
    subtitle.add_run("Executive Overview and User Journeys"),
    size=14,
    color="555555",
)

# Remove the original "At a glance" block and its small overview figure.
at_glance = next(i for i, p in enumerate(original) if p.text.strip() == "At a glance")
who_takes_part = next(
    i for i, p in enumerate(original) if p.text.strip() == "Who takes part?"
)
for p in original[at_glance:who_takes_part]:
    remove_paragraph(p)

# Remove the two detailed sequence figures and their captions.
for caption_text in (
    "Publisher journey through the Secured MCP Registry",
    "AI agent and LLM application journey",
):
    current = list(doc.paragraphs)
    caption_index = next(
        i for i, p in enumerate(current) if p.text.strip() == caption_text
    )
    caption = current[caption_index]
    drawing = current[caption_index - 1]
    remove_paragraph(drawing)
    remove_paragraph(caption)

# Build the new executive block at the end, then move it into the opening slot.
blocks = []
blocks.append(add_heading(doc, "Executive summary", 1))
blocks.append(
    add_body(
        doc,
        "SecureMCP provides a governed way for AI applications to find and use external capabilities. Publishers submit MCP services to the Secured MCP Registry, where identity, safety and permission information can be reviewed before those services are made available.",
    )
)
blocks.append(
    add_body(
        doc,
        "When an AI application needs a capability, its Agent Host and MCP Client use an approved configuration to connect to a Secured MCP Server. The server returns the available tool descriptions, the language model selects an appropriate tool, and the MCP Client carries the request and response between the application and the server.",
    )
)
blocks.append(
    add_body(
        doc,
        "SecureMCP applies safeguards throughout this journey. Policy and consent establish what is allowed; trust checks confirm the service and version; the Reflexive Core watches for changing behaviour; and provenance, audit and alerts keep important activity visible and accountable.",
    )
)
blocks.append(
    add_body(
        doc,
        "In simple terms: the registry decides what may be trusted, the MCP Client manages the connection, the Secured MCP Server performs the approved work, and SecureMCP continues to monitor the interaction.",
        lead="In simple terms:",
    )
)

blocks.append(add_heading(doc, "How data flows through SecureMCP", 1))
blocks.append(
    add_body(
        doc,
        "The operating model can be understood as five business steps. The detailed MCP protocol sequence is intentionally kept outside this executive view.",
    )
)
figure = doc.add_paragraph(style="Normal")
figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
figure.paragraph_format.space_before = Pt(6)
figure.paragraph_format.space_after = Pt(4)
figure.add_run().add_picture(str(FLOW), width=Inches(6.2))
blocks.append(figure)
caption = doc.add_paragraph(style="Normal")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.space_after = Pt(12)
set_arial(
    caption.add_run(
        "Five-step view: publish, verify, discover, use safely, and monitor"
    ),
    size=10,
    color="555555",
)
blocks.append(caption)

blocks.append(add_heading(doc, "What each part does", 1))
module_descriptions = [
    (
        "AI Agent / LLM Application",
        "understands the user’s request, sees the approved tools made available by the Host, and decides when a tool may help.",
    ),
    (
        "Agent Host",
        "runs the user experience and language model, creates MCP Client connections, applies approvals, and provides tool descriptions and results to the model.",
    ),
    (
        "MCP Client",
        "maintains one protocol connection to an MCP Server, requests the tool list, sends approved tool calls, and returns tool responses to the Host.",
    ),
    (
        "Secured MCP Registry",
        "is the governed system of record for publishers, MCP services, approved versions, connection information and lifecycle status.",
    ),
    (
        "Trust & Verification",
        "confirms the publisher, server version, certification, attestation and current suspension or revocation status.",
    ),
    (
        "Policy Enforcement",
        "checks whether a user, agent or MCP Client may perform a particular action under the organization’s rules.",
    ),
    (
        "Reflexive Core",
        "compares current activity with expected behaviour and can allow, restrict, require approval or stop an action as risk changes.",
    ),
    (
        "Secured MCP Server",
        "publishes tool descriptions and performs approved tool operations through SecureMCP runtime controls.",
    ),
    (
        "Provenance & Audit",
        "records important identities, decisions and outcomes so activity can be traced and reviewed.",
    ),
    (
        "Alerts & Notifications",
        "informs the responsible user, publisher or administrator when attention or action is required.",
    ),
]
for title, detail in module_descriptions:
    blocks.append(add_numbered_item(doc, title, detail))

blocks.append(
    add_body(
        doc,
        "Document scope: This is the business-level operating model. A separate technical appendix can show initialization, tools/list, tools/call, normal/elevated/critical branches, connection shutdown and transport-specific details.",
        lead="Document scope:",
    )
)

# Insert the new block immediately after the opening purpose paragraph.
anchor = original[2]._element
for block in blocks:
    anchor = move_after(block, anchor)

# Use explicit section-local numbering so Google Docs does not continue unrelated
# journeys at 11, 15, 21, and so on.
numbered_sections = {
    "What each part does",
    "Who takes part?",
    "Publisher journey",
    "AI Agent / LLM Application journey",
    "The message to take forward",
}
counter = None
for paragraph in list(doc.paragraphs):
    text = paragraph.text.strip()
    if text in numbered_sections:
        counter = 0
        continue
    if paragraph.style.name == "List Number":
        if not text:
            remove_paragraph(paragraph)
            continue
        if counter is None:
            counter = 0
        counter += 1
        paragraph.style = doc.styles["Normal"]
        paragraph.paragraph_format.left_indent = Inches(0.28)
        paragraph.paragraph_format.first_line_indent = Inches(-0.28)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        for run in paragraph.runs:
            set_arial(run, size=10.5)
        number_run = paragraph.add_run(f"{counter}. ")
        paragraph._p.remove(number_run._r)
        paragraph._p.insert(0, number_run._r)
        set_arial(number_run, size=10.5)
    elif paragraph.style.name.startswith("Heading"):
        counter = None

doc.core_properties.title = (
    "Secured MCP Registry - Executive Overview and User Journeys"
)
doc.save(OUTPUT)
print(OUTPUT)
