from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "mcp-registry-comparison-and-use-cases.docx"


def set_font(run, size=11, bold=False, color="000000"):
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Arial")
    rpr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=8, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_body(doc, text, *, lead=None):
    p = doc.add_paragraph()
    set_spacing(p)
    if lead and text.startswith(lead):
        set_font(p.add_run(lead), bold=True)
        set_font(p.add_run(text[len(lead) :]))
    else:
        set_font(p.add_run(text))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    sizes = {1: 20, 2: 16, 3: 14}
    before = {1: 20, 2: 18, 3: 16}
    after = {1: 6, 2: 6, 3: 4}
    set_spacing(p, before=before[level], after=after[level])
    set_font(p.add_run(text), size=sizes[level], color="000000")
    p.paragraph_format.keep_with_next = True
    return p


def add_number(doc, title, detail):
    p = doc.add_paragraph(style="List Number")
    set_spacing(p, after=6)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_font(p.add_run(title + " — "), bold=True)
    set_font(p.add_run(detail))
    return p


def shade_cell(cell, fill="F3F4F6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def fill_cell(cell, text, *, bold=False, size=9.5):
    p = cell.paragraphs[0]
    p.clear()
    set_spacing(p, after=0, line=1.08)
    set_font(p.add_run(text), size=size, bold=bold)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

for level, size, before, after, color in [
    (1, 20, 20, 6, "000000"),
    (2, 16, 18, 6, "000000"),
    (3, 14, 16, 4, "434343"),
]:
    style = doc.styles[f"Heading {level}"]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(size)
    style.font.bold = False
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

title = doc.add_paragraph()
set_spacing(title, after=3)
set_font(title.add_run("MCP Registry Landscape"), size=26)

subtitle = doc.add_paragraph()
set_spacing(subtitle, after=18)
set_font(
    subtitle.add_run("Comparison, SecureMCP Positioning and Business Use Cases"),
    size=14,
    color="555555",
)

add_body(
    doc,
    "A non-technical comparison of how Claude Code, Codex and other MCP ecosystems discover and connect to MCP services, followed by practical use cases for a Secured MCP Registry.",
)

add_heading(doc, "Executive conclusion", 1)
add_number(
    doc,
    "Claude Code",
    "connects to MCP servers that a user or team configures. Its documentation provides examples and points users toward community servers, but it does not describe Claude Code as automatically selecting servers from one mandatory registry.",
)
add_number(
    doc,
    "Codex",
    "also connects to explicitly configured MCP servers. Local Codex clients share MCP configuration, while hosted ChatGPT and Codex experiences can receive MCP-backed tools through plugins.",
)
add_number(
    doc,
    "Public registries",
    "primarily help people discover, publish, install or host MCP servers. Their focus ranges from open metadata to managed credentials, containers, gateways and operational monitoring.",
)
add_number(
    doc,
    "Secured MCP Registry",
    "is positioned as a governed control plane rather than only a directory. It combines publication and discovery with certification, policy, contracts, consent, provenance, adaptive risk controls and lifecycle actions.",
)

add_heading(doc, "Important terminology", 1)
add_body(
    doc,
    "Claude Code and Codex are MCP clients, not registries. A registry helps users discover or govern MCP servers; the client connects to the selected server after configuration or installation.",
)
add_body(
    doc,
    "Therefore, the most accurate question is not “Which registry does Claude Code or Codex use?” but “How do these clients discover, configure and govern MCP servers?”",
)

add_heading(doc, "How Claude Code works with MCP", 1)
add_number(
    doc,
    "Connection model",
    "Users add local or remote MCP servers through Claude Code commands or configuration.",
)
add_number(
    doc,
    "Team sharing",
    "Project-scoped server definitions can be stored in a shared .mcp.json file.",
)
add_number(
    doc,
    "User approval",
    "Claude Code asks for approval before using project-scoped MCP servers from shared configuration.",
)
add_number(
    doc,
    "Authentication",
    "Remote HTTP and SSE servers can use OAuth, headers or other configured credentials.",
)
add_number(
    doc,
    "Discovery",
    "Anthropic documentation lists popular servers and directs users to broader community sources, while installation remains an explicit user or team action.",
)
add_number(
    doc,
    "Security message",
    "Anthropic warns that third-party servers are not all verified and should be installed only when trusted.",
)

add_heading(doc, "How Codex works with MCP", 1)
add_number(
    doc,
    "Connection model",
    "Codex supports local STDIO servers and remote Streamable HTTP servers.",
)
add_number(
    doc,
    "Shared configuration",
    "The ChatGPT desktop app, Codex CLI and Codex IDE extension share MCP configuration on the same Codex host.",
)
add_number(
    doc,
    "Project control",
    "Servers can be configured globally or in a trusted project through config.toml.",
)
add_number(
    doc,
    "Tool control",
    "Codex supports server enablement, required-server settings, tool allowlists, tool denylists and approval modes.",
)
add_number(
    doc,
    "Hosted tools",
    "ChatGPT and Codex can also receive remote MCP-backed tools through installed plugins rather than local MCP configuration.",
)
add_number(
    doc,
    "Discovery",
    "Official Codex documentation provides examples of useful MCP servers, but does not describe a built-in public-registry search command comparable to GitHub Copilot CLI’s registry search.",
)

add_heading(doc, "Registry and catalogue comparison", 1)

comparisons = [
    (
        "1. Official MCP Registry",
        "A community-driven central repository and API for publishing and discovering MCP server metadata.",
        "Open ecosystem metadata, standard registry interfaces and interoperability.",
        "It is currently described as preview. Published version metadata is immutable, and its FAQ directs malicious-server reports to the underlying package registry and a GitHub issue.",
        "Best used as an ecosystem source of record or upstream feed, not by itself as a complete enterprise runtime-security system.",
    ),
    (
        "2. GitHub MCP Registry",
        "A discovery experience for finding official or relevant MCP servers within GitHub-oriented developer workflows.",
        "Fast discovery and installation, plus enterprise support for internal registries and allowlist controls in supported GitHub environments.",
        "Its strongest value is developer workflow integration and organizational distribution policy.",
        "A useful comparison for self-service discovery, but SecureMCP goes further into execution-time governance and adaptive risk.",
    ),
    (
        "3. Docker MCP Catalog and Toolkit",
        "A curated catalogue designed to make MCP servers easy to discover, deploy and run with Docker tooling.",
        "Container packaging, repeatable deployment and desktop-based operational convenience.",
        "It is especially strong when the main problem is safely packaging and running third-party server software.",
        "SecureMCP can complement this model by governing which containerized server may be used, by whom and under which policy.",
    ),
    (
        "4. Smithery",
        "An open registry combined with managed MCP connections, OAuth, token refresh and credential storage.",
        "Low-friction connection management for agent builders and multi-user applications.",
        "It reduces integration effort by managing authentication and connection lifecycle.",
        "SecureMCP differentiates through organization-owned policy, certification, consent, contracts, provenance and reflexive controls.",
    ),
    (
        "5. Glama",
        "A large MCP directory with hosting, deployment, gateway, logs, analytics and per-tool access controls.",
        "Discovery plus production operations, health checks, hosting and connection-profile security.",
        "It combines a public directory with operational infrastructure and runtime visibility.",
        "It is the closest operational comparison, while SecureMCP emphasizes an integrated governance model across the full trust lifecycle.",
    ),
    (
        "6. PureCipher Secured MCP Registry",
        "A governed registry and security control plane built around SecureMCP.",
        "Publisher verification, certification, policy, contracts, consent, provenance, Reflexive Core, alerts, moderation and lifecycle control.",
        "It is intended to make trust decisions before publication, during discovery and while an MCP capability is running.",
        "Its clearest position is “governed MCP adoption for organizations,” rather than “the largest public directory.”",
    ),
]

for heading, role, strength, boundary, position in comparisons:
    add_heading(doc, heading, 2)
    add_body(doc, f"Primary role: {role}", lead="Primary role:")
    add_body(doc, f"Strength: {strength}", lead="Strength:")
    add_body(doc, f"Boundary: {boundary}", lead="Boundary:")
    add_body(doc, f"SecureMCP relevance: {position}", lead="SecureMCP relevance:")

add_heading(doc, "Comparison at a glance", 1)
headers = ["Option", "Best suited to", "Discovery model", "Governance emphasis"]
rows = [
    [
        "Official MCP Registry",
        "Open ecosystem metadata",
        "Public API and registry",
        "Publisher metadata and ecosystem reporting",
    ],
    [
        "GitHub MCP Registry",
        "Developer self-service",
        "GitHub-integrated search",
        "Enterprise allowlists in supported GitHub environments",
    ],
    [
        "Docker MCP Catalog",
        "Packaged deployment",
        "Curated Docker catalogue",
        "Container execution and operational consistency",
    ],
    [
        "Smithery",
        "Managed connections",
        "Open registry and connection API",
        "OAuth, credentials and session lifecycle",
    ],
    [
        "Glama",
        "Hosted MCP operations",
        "Directory plus deployment",
        "Gateway, logs, health and per-tool controls",
    ],
    [
        "Secured MCP Registry",
        "Governed organizational adoption",
        "Approved internal/public catalogue",
        "Policy, consent, contracts, provenance and adaptive risk",
    ],
]

table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
for i, text in enumerate(headers):
    fill_cell(table.rows[0].cells[i], text, bold=True, size=9)
    shade_cell(table.rows[0].cells[i], "F3F4F6")
for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        fill_cell(cells[i], text, size=8.7)
set_table_geometry(table, [1800, 2340, 2160, 3060])

add_heading(doc, "Where the Secured MCP Registry is different", 1)
differences = [
    (
        "Governance before discovery",
        "Only approved publishers, services and versions should become visible to the intended audience.",
    ),
    (
        "Policy during execution",
        "A service can be discoverable but still restricted for a particular user, agent, purpose or environment.",
    ),
    (
        "Consent and contracts",
        "Access can depend on explicit permission and an agreed purpose rather than on installation alone.",
    ),
    (
        "Traceable provenance",
        "Important requests, decisions and outcomes can be recorded for investigation and accountability.",
    ),
    (
        "Adaptive protection",
        "The Reflexive Core can detect behavioural drift and escalate from normal operation to approval, restriction or blocking.",
    ),
    (
        "Lifecycle enforcement",
        "Moderators can approve, suspend, deprecate, revoke or deregister services and versions.",
    ),
    (
        "Client identity",
        "AI agents, services and frameworks can receive controlled identities and tokens instead of appearing as anonymous consumers.",
    ),
    (
        "Enterprise deployment choice",
        "The registry can be operated as an internal catalogue, a curated public service or a federated trust environment.",
    ),
]
for name, detail in differences:
    add_number(doc, name, detail)

add_heading(doc, "Business use cases", 1)
use_cases = [
    (
        "Enterprise-approved AI tools",
        "Employees want Claude Code, Codex or another AI client to use Jira, GitHub, databases and internal APIs.",
        "The organization publishes an approved catalogue and blocks or discourages unreviewed MCP servers.",
        "Registry, identity, certification, policy and client tokens.",
    ),
    (
        "Regulated financial-services agent",
        "An AI agent prepares a customer or transaction analysis using sensitive systems.",
        "Consent, purpose limits and audit evidence ensure that the agent accesses only approved data for the approved task.",
        "Policy, contracts, consent, provenance, compliance and alerts.",
    ),
    (
        "Healthcare workflow assistant",
        "A clinical or administrative assistant needs controlled access to patient-related services.",
        "The registry approves specific service versions, while runtime controls limit access by role, purpose and environment.",
        "Certification, policy, consent, sandboxing, provenance and revocation.",
    ),
    (
        "Third-party MCP vendor onboarding",
        "A SaaS vendor wants its MCP service to be available to enterprise AI agents.",
        "The publisher submits identity, capability, permission and version evidence; reviewers approve or request remediation.",
        "Publisher service, certification, moderation, registry and notifications.",
    ),
    (
        "Developer self-service with guardrails",
        "Engineering teams want easy access to documentation, observability, browser and repository tools.",
        "Developers discover pre-approved servers without waiting for a manual installation review each time.",
        "Catalogue, install guidance, policy profiles and client identity.",
    ),
    (
        "Secure API-to-MCP conversion",
        "An organization wants to expose selected internal REST operations to AI agents.",
        "OpenAPI operations are converted into governed MCP toolsets while credentials remain controlled.",
        "OpenAPI ingestion, gateway, policy, contracts, provenance and alerts.",
    ),
    (
        "High-risk production action",
        "An AI agent requests a deployment, payment, account change or destructive infrastructure action.",
        "The request is checked against policy and consent, executed in a restricted environment where appropriate and stopped when behaviour becomes abnormal.",
        "Policy, consent, contracts, sandbox, Reflexive Core and audit.",
    ),
    (
        "Rapid incident response",
        "A trusted server version or publisher credential becomes compromised.",
        "Administrators revoke trust, suspend the listing, notify affected clients and prevent new connections.",
        "Revocation, registry lifecycle, alerts, client inventory and provenance.",
    ),
    (
        "Multi-agent workflow control",
        "A coordinator agent delegates work to specialist agents that use different MCP services.",
        "Each agent receives a distinct identity and only the tools required for its assigned role.",
        "Client identity, policy, contracts, consent and provenance.",
    ),
    (
        "Federated organization",
        "Business units or partner organizations want to share approved MCP services without surrendering local governance.",
        "Trust information can be exchanged while each organization keeps its own approval and policy decisions.",
        "Federation, registry, certification, revocation and local policy.",
    ),
    (
        "Controlled version promotion",
        "A publisher needs to move an MCP service from development to testing and production.",
        "Each version is reviewed, certified and promoted according to environment-specific policy.",
        "Policy versioning, certification, moderation, provenance and notifications.",
    ),
    (
        "Shadow MCP visibility",
        "Teams are using untracked MCP servers with unknown permissions or credentials.",
        "The organization establishes a registered client inventory and an approved catalogue, making exceptions visible and reviewable.",
        "Client registry, catalogue, policy, audit and alerts.",
    ),
]

for index, (name, situation, response, pillars) in enumerate(use_cases, start=1):
    add_heading(doc, f"{index}. {name}", 2)
    add_body(doc, f"Situation: {situation}", lead="Situation:")
    add_body(doc, f"Registry response: {response}", lead="Registry response:")
    add_body(doc, f"Key pillars: {pillars}", lead="Key pillars:")

add_heading(doc, "Recommended positioning", 1)
add_number(
    doc,
    "Do not compete on catalogue size alone",
    "Public registries and directories are already optimized for broad discovery.",
)
add_number(
    doc,
    "Use public registries as upstream sources",
    "Import or reference ecosystem metadata, then apply organizational approval and security enrichment.",
)
add_number(
    doc,
    "Lead with governed adoption",
    "Position the Secured MCP Registry as the place where an organization decides which MCP services and versions may be trusted.",
)
add_number(
    doc,
    "Emphasize runtime protection",
    "The strongest distinction is the combination of policy, consent, contracts, provenance and the Reflexive Core after installation.",
)
add_number(
    doc,
    "Support multiple clients",
    "The registry should serve Claude Code, Codex and any standards-compliant MCP client rather than becoming tied to one model vendor.",
)

add_heading(doc, "Sources", 1)
sources = [
    (
        "Anthropic: Connect Claude Code to tools via MCP",
        "https://docs.anthropic.com/en/docs/claude-code/mcp",
    ),
    (
        "OpenAI Codex Manual: Model Context Protocol",
        "https://learn.chatgpt.com/docs/extend/mcp",
    ),
    ("Official MCP Registry API", "https://registry.modelcontextprotocol.io/docs"),
    ("Official MCP Registry FAQ", "https://modelcontextprotocol.io/registry/faq"),
    (
        "GitHub: Meet the GitHub MCP Registry",
        "https://github.blog/ai-and-ml/github-copilot/meet-the-github-mcp-registry-the-fastest-way-to-discover-mcp-servers/",
    ),
    (
        "GitHub: Internal MCP registry and allowlist controls",
        "https://github.blog/changelog/2025-09-12-internal-mcp-registry-and-allowlist-controls-for-vs-code-insiders/",
    ),
    ("Docker MCP Registry", "https://github.com/docker/mcp-registry"),
    ("Smithery documentation", "https://smithery.ai/docs"),
    ("Smithery Connect", "https://smithery.ai/docs/use/connect"),
    ("Glama MCP hosting and gateway", "https://glama.ai/mcp/hosting"),
]
for name, url in sources:
    add_number(doc, name, url)

add_body(
    doc,
    "Research note: Public product capabilities change quickly. This comparison reflects official documentation reviewed on 28 July 2026. Statements about the PureCipher Secured MCP Registry are based on the current project implementation and architecture.",
)

props = doc.core_properties
props.title = "MCP Registry Landscape: Comparison and Use Cases"
props.subject = "Comparison of MCP registries and SecureMCP business use cases"
props.author = "PureCipher"

doc.save(OUTPUT)
print(OUTPUT)
