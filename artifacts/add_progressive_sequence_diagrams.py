from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path("/Users/purecipher/code/xsecuremcp2.0/artifacts")
SOURCE = ROOT / "Secured MCP Registry - Executive Overview.docx"
OUTPUT = ROOT / "Secured MCP Registry - Progressive Technical Guide.docx"

PUBLISHER = ROOT / "publisher-journey-sequence.png"
AGENT = ROOT / "ai-agent-journey-sequence.png"
PROTOCOL = ROOT / "mcp-client-server-protocol-sequence.png"
RUNTIME = ROOT / "secured-mcp-reflexive-sequence.png"
COMPARISON = Path(
    "/Users/purecipher/Downloads/mcp-registry-comparison-and-use-cases-google-docs.docx"
)


def set_arial(run, size=11, bold=False, color="000000"):
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Arial")
    rpr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    sizes = {1: 20, 2: 16, 3: 13}
    set_arial(p.add_run(text), sizes[level])
    return p


def body(doc, text, lead=None):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if lead and text.startswith(lead):
        set_arial(p.add_run(lead), 10.5, True)
        set_arial(p.add_run(text[len(lead) :]), 10.5)
    else:
        set_arial(p.add_run(text), 10.5)
    return p


def numbered(doc, number, title, detail):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.28)
    p.paragraph_format.space_after = Pt(3)
    set_arial(p.add_run(f"{number}. "), 10.5)
    set_arial(p.add_run(f"{title} — "), 10.5, True)
    set_arial(p.add_run(detail), 10.5)
    return p


def figure(doc, image, caption, width=None, height=None):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    kwargs = {}
    if width is not None:
        kwargs["width"] = Inches(width)
    if height is not None:
        kwargs["height"] = Inches(height)
    p.add_run().add_picture(str(image), **kwargs)
    c = doc.add_paragraph(style="Normal")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(4)
    set_arial(c.add_run(caption), 9, False, "555555")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def add_comparison_table(doc, source_table):
    widths = [2100, 2250, 2100, 2910]
    table = doc.add_table(rows=0, cols=4)
    table.autofit = False
    table.style = "Table Grid"
    for row_index, source_row in enumerate(source_table.rows):
        cells = table.add_row().cells
        for column_index, cell in enumerate(cells):
            set_cell_width(cell, widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = source_row.cells[column_index].text
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_arial(run, 8.25, row_index == 0)
            if row_index == 0:
                shade_cell(cell, "EDE9FE")
        if row_index == 0:
            tr_pr = cells[0]._tc.getparent().get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
    table.rows[0].cells[0]._tc.getparent()
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(8)
    return table


doc = Document(SOURCE)

# Expand the subtitle without changing the established first-page design.
subtitle = doc.paragraphs[1]
subtitle.clear()
set_arial(
    subtitle.add_run("Executive Overview, User Journeys and Technical Sequences"),
    14,
    False,
    "555555",
)

# The appended walkthrough uses landscape pages so sequence labels remain readable.
section = doc.add_section(WD_SECTION.NEW_PAGE)
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(11)
section.page_height = Inches(8.5)
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

heading(doc, "Progressive technical walkthrough", 1)
body(
    doc,
    "The diagrams below move from user intent to implementation detail. Each level adds participants and decisions while preserving the same operating model.",
)
numbered(
    doc,
    1,
    "Level 1: Executive flow",
    "the five-step diagram near the beginning explains publish, verify, discover, use safely, and monitor.",
)
numbered(
    doc,
    2,
    "Level 2: User journeys",
    "separate publisher and AI-agent sequences show who initiates each action and what outcome is returned.",
)
numbered(
    doc,
    3,
    "Level 3: MCP protocol",
    "the client/server sequence shows initialization, tool discovery, tool calls, responses, and shutdown.",
)
numbered(
    doc,
    4,
    "Level 4: SecureMCP runtime controls",
    "the final sequence adds trust, policy, Reflexive Core, provenance, audit, and alert decisions.",
)
body(
    doc,
    "Reading convention: solid arrows are requests or actions; dashed arrows are returned decisions or results; alternate frames show different outcomes.",
    lead="Reading convention:",
)

doc.add_page_break()
heading(doc, "Level 2A — Publisher sequence", 1)
body(
    doc,
    "This view follows an MCP service from submission through review, publication, updates, suspension, or retirement.",
)
figure(
    doc,
    PUBLISHER,
    "Publisher submission, approval and lifecycle sequence",
    height=5.25,
)

doc.add_page_break()
heading(doc, "Level 2B — AI Agent / LLM Application sequence", 1)
body(
    doc,
    "This view follows a user request from capability discovery through governed execution. It introduces normal, elevated-risk, and critical-risk outcomes without exposing every internal service.",
)
figure(
    doc,
    AGENT,
    "AI-agent discovery, authorization, execution and risk-response sequence",
    height=5.15,
)

doc.add_page_break()
heading(doc, "Level 3 — Standard MCP client/server sequence", 1)
body(
    doc,
    "This view separates the Agent Host, MCP Client, language model, and MCP Server. It also shows where the client and server start and stop.",
)
numbered(
    doc,
    1,
    "Client lifecycle",
    "the Agent Host starts and closes the MCP Client.",
)
numbered(
    doc,
    2,
    "Server lifecycle",
    "a local server may be started and stopped with the client; a remote server normally remains running while the transport connection closes.",
)
figure(
    doc,
    PROTOCOL,
    "MCP initialization, tools/list, tools/call, response and shutdown sequence",
    height=4.95,
)

doc.add_page_break()
heading(doc, "Level 4 — SecureMCP runtime control sequence", 1)
body(
    doc,
    "This most technical view expands the governed runtime. The Registry verifies the approved service and configuration; Policy Enforcement decides whether the requested action is allowed; the Reflexive Core evaluates current behaviour and changing risk; Provenance and Audit record the decision; and Alerts and Notifications surface events requiring attention.",
)
numbered(
    doc,
    1,
    "Normal behaviour",
    "execution continues and the result is returned.",
)
numbered(
    doc,
    2,
    "Elevated risk",
    "execution may require approval, be restricted, or return a limited response.",
)
numbered(
    doc,
    3,
    "Critical risk",
    "execution is blocked, security status may be updated, and responsible owners are notified.",
)
figure(
    doc,
    RUNTIME,
    "SecureMCP verification, policy, Reflexive Core, audit and alert sequence",
    height=4.55,
)

# Return to portrait for the merged landscape comparison and business-use-case content.
section = doc.add_section(WD_SECTION.NEW_PAGE)
section.orientation = WD_ORIENT.PORTRAIT
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

heading(doc, "Registry comparison and business use cases", 1)
body(
    doc,
    "This section compares how common MCP ecosystems support discovery and connection, then applies the Secured MCP Registry to practical organizational scenarios.",
)

comparison_doc = Document(COMPARISON)
comparison_table_added = False
list_counter = 0
for source_paragraph in comparison_doc.paragraphs[3:]:
    text = source_paragraph.text.strip()
    if not text:
        continue
    style_name = source_paragraph.style.name
    if style_name == "Heading 1":
        heading(doc, text, 1)
        list_counter = 0
        if text == "Comparison at a glance" and not comparison_table_added:
            add_comparison_table(doc, comparison_doc.tables[0])
            comparison_table_added = True
    elif style_name == "Heading 2":
        heading(doc, text, 2)
        list_counter = 0
    elif style_name == "List Number":
        list_counter += 1
        if " — " in text:
            title, detail = text.split(" — ", 1)
        else:
            title, detail = text, ""
        numbered(doc, list_counter, title, detail)
    else:
        lead = None
        for candidate in (
            "Primary role:",
            "Strength:",
            "Boundary:",
            "SecureMCP relevance:",
            "Situation:",
            "Registry response:",
            "Key pillars:",
            "Research note:",
        ):
            if text.startswith(candidate):
                lead = candidate
                break
        body(doc, text, lead=lead)

doc.core_properties.title = "Secured MCP Registry - Progressive Technical Guide"
doc.save(OUTPUT)
print(OUTPUT)
