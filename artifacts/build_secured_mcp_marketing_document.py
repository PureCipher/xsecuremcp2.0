from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/purecipher/code/xsecuremcp2.0/artifacts")
OUTPUT = ROOT / "Secured MCP Registry - Marketing Overview.docx"
FLOW = ROOT / "secured-mcp-high-level-flow.png"

NAVY = "18233A"
PURPLE = "6D5CE7"
PURPLE_DARK = "5142C7"
LAVENDER = "F0EEFF"
PALE = "F6F7FB"
GRAY = "5F6673"
LIGHT_GRAY = "D9DDE7"
WHITE = "FFFFFF"
BLACK = "111318"


def font(run, size=11, bold=False, color=BLACK, italic=False):
    run.font.name = "Aptos"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Aptos")
    rpr.rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=120, start=150, bottom=120, end=150):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")


def paragraph(doc, text="", size=11, bold=False, color=BLACK, after=7, before=0,
              align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    font(p.add_run(text), size, bold, color, italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    sizes = {1: 19, 2: 14, 3: 11.5}
    before = {1: 12, 2: 10, 3: 7}
    after = {1: 7, 2: 5, 3: 3}
    colors = {1: NAVY, 2: PURPLE_DARK, 3: NAVY}
    p.paragraph_format.space_before = Pt(before[level])
    p.paragraph_format.space_after = Pt(after[level])
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text), sizes[level], level == 3, colors[level])
    return p


def label_body(doc, label, text, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.16
    font(p.add_run(label), 10.5, True, NAVY)
    font(p.add_run(text), 10.5, False, BLACK)
    return p


def kicker(doc, text):
    p = paragraph(doc, text.upper(), 9.5, True, PURPLE, after=5)
    p.paragraph_format.keep_with_next = True
    return p


def callout(doc, title, text, fill=LAVENDER):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_width(cell, 9120)
    cell_margins(cell, 160, 190, 160, 190)
    shade(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    font(p.add_run(title), 12, True, PURPLE_DARK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    font(p2.add_run(text), 10.5, False, NAVY)
    paragraph(doc, "", after=3)


def page_break(doc):
    doc.add_page_break()


def feature(doc, number, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(f"{number}. "), 12.5, True, PURPLE)
    font(p.add_run(title), 12.5, True, NAVY)
    paragraph(doc, text, 10.5, False, BLACK, after=7)


def use_case(doc, number, title, situation, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(f"{number}. "), 13, True, PURPLE)
    font(p.add_run(title), 13, True, NAVY)
    label_body(doc, "Business need: ", situation, after=3)
    label_body(doc, "SecureMCP value: ", value, after=6)


def comparison_table(doc):
    rows = [
        ("Capability", "Typical registry or catalogue", "Secured MCP Registry"),
        ("Discovery", "Find and install available MCP servers", "Discover approved services for the right audience"),
        ("Trust", "Publisher and package metadata", "Identity, certification, attestation and version status"),
        ("Access", "Configured by the user or client", "Policy, consent, contracts and client identity"),
        ("Runtime", "Connection and operational visibility", "Continuous behavioural risk assessment and adaptive controls"),
        ("Lifecycle", "Publish, update or remove listings", "Approve, suspend, revoke, deprecate and notify"),
        ("Evidence", "Logs vary by platform", "Provenance, audit and compliance-ready records"),
    ]
    widths = [1700, 3400, 4020]
    table = doc.add_table(rows=0, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, value in enumerate(row):
            set_cell_width(cells[ci], widths[ci])
            cell_margins(cells[ci], 95, 115, 95, 115)
            cells[ci].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri == 0:
                shade(cells[ci], PURPLE)
            elif ci == 2:
                shade(cells[ci], LAVENDER)
            p = cells[ci].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            font(p.add_run(value), 8.6, ri == 0 or ci == 0, WHITE if ri == 0 else BLACK)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.18

for level in (1, 2, 3):
    style = doc.styles[f"Heading {level}"]
    style.font.name = "Aptos"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")

# Quiet running header and footer.
header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(header.add_run("PURECIPHER  |  SECUREMCP"), 8.5, True, GRAY)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(footer.add_run("Secured MCP Registry  •  Marketing Overview"), 8, False, GRAY)

# Page 1 — editorial cover.
paragraph(doc, "", after=68)
kicker(doc, "Trusted AI Extensibility")
paragraph(doc, "Secured MCP Registry", 30, True, NAVY, after=8)
paragraph(
    doc,
    "Governed discovery. Trusted execution. Continuous assurance.",
    17,
    False,
    PURPLE_DARK,
    after=22,
)
paragraph(
    doc,
    "A security and governance control plane for publishing, discovering and using MCP services across AI agents and LLM-powered applications.",
    13,
    False,
    NAVY,
    after=28,
)
callout(
    doc,
    "Move from fragmented MCP adoption to governed innovation.",
    "SecureMCP helps organizations decide which publishers, services and versions may be trusted—then continues to apply policy, monitor behaviour and preserve accountability while those capabilities are used.",
)
paragraph(doc, "", after=70)
paragraph(doc, "Marketing Overview  |  July 2026", 10, True, GRAY, after=3)
paragraph(doc, "PureCipher", 10, False, GRAY, after=0)

# Page 2 — market problem and promise.
page_break(doc)
kicker(doc, "The opportunity")
heading(doc, "AI agents are becoming more capable. Their connections must become more governable.", 1)
paragraph(
    doc,
    "Model Context Protocol makes it easier for AI applications to reach tools, APIs and enterprise systems. That speed creates value—but it also introduces a new governance question: which capabilities should an AI agent be allowed to discover and use, under what conditions, and with what evidence?",
    11,
    after=11,
)
heading(doc, "The risk of unmanaged MCP adoption", 2)
feature(doc, 1, "Unknown trust", "A useful server can still have an unclear publisher, unreviewed version or excessive permissions.")
feature(doc, 2, "Fragmented control", "Teams can configure MCP services independently, creating inconsistent approval and security practices.")
feature(doc, 3, "Static decisions", "A service that was safe at installation may become risky as context, behaviour or credentials change.")
feature(doc, 4, "Limited accountability", "Without common provenance and audit records, important actions are harder to explain and investigate.")
callout(
    doc,
    "SecureMCP turns trust into an operating model.",
    "It combines a governed registry with verification, policy, consent, runtime protection, lifecycle control and auditable evidence.",
    PALE,
)

# Page 3 — five-step visual.
page_break(doc)
kicker(doc, "How it works")
heading(doc, "A simple path from publication to continuous assurance", 1)
paragraph(
    doc,
    "SecureMCP creates one governed journey for publishers, AI applications and security teams.",
    11,
    after=8,
)
pic = doc.add_paragraph()
pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
pic.paragraph_format.space_before = Pt(6)
pic.paragraph_format.space_after = Pt(8)
pic.add_run().add_picture(str(FLOW), width=Inches(6.55))
feature(doc, 1, "Publish", "A publisher submits an MCP service with ownership, purpose, permissions and version information.")
feature(doc, 2, "Verify and approve", "Identity, safety evidence, certification and organizational requirements are checked.")
feature(doc, 3, "Discover", "Eligible AI agents and applications find approved capabilities through a trusted catalogue.")
feature(doc, 4, "Use safely", "The MCP Client connects to the approved server while policy and consent govern the requested action.")
feature(doc, 5, "Monitor and adapt", "SecureMCP records activity and can allow, restrict, require approval or stop access as risk changes.")

# Page 4 — differentiators.
page_break(doc)
kicker(doc, "Why SecureMCP")
heading(doc, "More than a directory: a governed control plane for MCP", 1)
feature(doc, 1, "Trust before discovery", "Only approved publishers, services and versions become visible to the intended users and agents.")
feature(doc, 2, "Policy during execution", "A discoverable service can still be limited by user, agent, purpose, environment, data sensitivity or action risk.")
feature(doc, 3, "Consent and contracts", "Access reflects explicit authority and an agreed purpose—not installation alone.")
feature(doc, 4, "Reflexive protection", "The Reflexive Core compares current activity with expected behaviour and adapts the response when risk changes.")
feature(doc, 5, "Traceable accountability", "Provenance, audit, alerts and compliance evidence preserve who requested what, why a decision was made and what happened.")
feature(doc, 6, "Lifecycle enforcement", "Organizations can approve, suspend, revoke, deprecate or retire services and versions across their lifecycle.")
heading(doc, "The business result", 2)
callout(
    doc,
    "Faster adoption with clearer control",
    "Developers and business teams gain access to approved capabilities without repeating the same security review, while risk and governance teams retain a strategic control point.",
)

# Page 5 — business outcomes and first use cases.
page_break(doc)
kicker(doc, "Business value")
heading(doc, "What organizations gain", 1)
feature(doc, 1, "Accelerated innovation", "Publishers and internal teams can bring trusted MCP capabilities to users faster.")
feature(doc, 2, "Reduced integration risk", "Approved versions, permissions and identities replace ad hoc connections and unknown dependencies.")
feature(doc, 3, "Consistent governance", "One operating model can serve Claude Code, Codex and other standards-compliant MCP clients.")
feature(doc, 4, "Audit-ready evidence", "Important identities, decisions and outcomes remain available for review, investigation and assurance.")
heading(doc, "Representative use cases", 1)
use_case(doc, 1, "Enterprise-approved AI tools", "Employees want AI clients to use Jira, GitHub, databases and internal APIs.", "Provide a curated catalogue of reviewed services and discourage or block unapproved alternatives.")
use_case(doc, 2, "Regulated financial-services agent", "An AI agent analyzes customers or transactions using sensitive systems.", "Apply purpose limits, consent, policy and audit evidence so access matches the approved task.")
use_case(doc, 3, "Healthcare workflow assistant", "A clinical or administrative assistant needs controlled access to patient-related services.", "Approve specific service versions and limit access by role, purpose and environment.")

# Page 6 — additional use cases and position.
page_break(doc)
kicker(doc, "Secure adoption at scale")
heading(doc, "Use cases across the MCP lifecycle", 1)
use_case(doc, 4, "Third-party vendor onboarding", "A SaaS provider wants its MCP service to reach enterprise AI agents.", "Review publisher identity, capability, permissions and version evidence before publication.")
use_case(doc, 5, "High-risk production action", "An agent requests a deployment, payment, account change or destructive infrastructure action.", "Check policy and consent, use restricted execution where appropriate and stop abnormal behaviour.")
use_case(doc, 6, "Rapid incident response", "A trusted version or publisher credential becomes compromised.", "Revoke trust, suspend the listing, notify affected clients and prevent new connections.")
heading(doc, "From registry to control plane", 1)
comparison_table(doc)

# Page 7 — deployment and close.
page_break(doc)
kicker(doc, "Designed for the way organizations operate")
heading(doc, "One trust model. Multiple deployment choices.", 1)
feature(doc, 1, "Internal registry", "Create a private catalogue of approved MCP services for employees, developers and enterprise AI agents.")
feature(doc, 2, "Curated public service", "Publish approved services to a controlled audience while maintaining certification and lifecycle status.")
feature(doc, 3, "Federated trust environment", "Share approved trust information across business units or partners while preserving local policy decisions.")
heading(doc, "Works across the MCP ecosystem", 2)
paragraph(
    doc,
    "The Secured MCP Registry is designed to complement—not replace—the MCP standard and public ecosystem. Organizations can use public registries and catalogues as upstream sources, enrich that metadata with their own approval and security evidence, and serve multiple standards-compliant clients through one governed model.",
    11,
    after=12,
)
callout(
    doc,
    "The positioning is clear",
    "Public directories optimize for broad discovery. SecureMCP optimizes for governed organizational adoption—before publication, during discovery and throughout execution.",
)
heading(doc, "Start with a focused adoption path", 1)
feature(doc, 1, "Choose a priority workflow", "Select one valuable MCP-enabled use case with clear users, systems and business outcomes.")
feature(doc, 2, "Define the trust requirements", "Identify the publishers, versions, permissions, policies, consent and evidence required.")
feature(doc, 3, "Publish an approved catalogue", "Make trusted capabilities easy for eligible AI clients to discover and use.")
feature(doc, 4, "Expand with evidence", "Use runtime insight, audit records and operating experience to scale adoption responsibly.")
paragraph(doc, "SecureMCP", 18, True, PURPLE_DARK, after=2)
paragraph(doc, "Trust the capability. Govern the action. Preserve the evidence.", 12, True, NAVY, after=0)

doc.core_properties.title = "Secured MCP Registry - Marketing Overview"
doc.core_properties.subject = "Marketing overview for the PureCipher Secured MCP Registry"
doc.core_properties.author = "PureCipher"
doc.save(OUTPUT)
print(OUTPUT)
